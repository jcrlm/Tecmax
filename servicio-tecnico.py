#!/usr/bin/env python
# ============================================================
# TECMAX - Sistema de Servicio Técnico y Facturación
# Técnico: Camilo Restrepo
# Plataforma: Termux / Android
# ============================================================

from docx import Document
from docx.shared import Inches
from datetime import datetime
import os, json, subprocess
from reportlab.lib.pagesizes import A5
from reportlab.pdfgen import canvas

# ================= CONFIGURACIÓN =================

EMPRESA = "TECMAX"
TECNICO = "Camilo Restrepo"

BASE = os.path.expanduser("~/storage/shared/tecmax")
BASE_CLIENTES = f"{BASE}/clientes"
BASE_INSUMOS = f"{BASE}/insumos"
LOGO = f"{BASE}/logo.png"

CONTADOR = os.path.expanduser("~/bin/contador_facturas.txt")

os.makedirs(BASE_CLIENTES, exist_ok=True)
os.makedirs(BASE_INSUMOS, exist_ok=True)

# ================= FUNCIONES =================

def pedir_texto(msg):
    while True:
        t = input(msg).strip()
        if t:
            return t
        print("❌ No puede estar vacío.\n")

def pedir_precio(msg):
    while True:
        txt = input(msg)
        try:
            return float(
                txt.lower()
                .replace("$","")
                .replace("mil","000")
                .replace(".","")
                .replace(",","")
                .strip()
            )
        except ValueError:
            print("❌ Precio inválido. Ej: 20000 | 20 mil\n")

def pedir_opcion(titulo, opciones):
    while True:
        print(titulo)
        for k,v in opciones.items():
            print(f"{k}) {v}")
        op = input("Seleccione: ").strip()
        if op in opciones:
            return opciones[op]
        print("❌ Opción inválida\n")

def slug(t):
    return t.replace(" ","_")

def enviar_whatsapp(numero, mensaje):
    numero = numero.replace(" ","").replace("-","")
    url = f"https://wa.me/{numero}?text={mensaje.replace(' ','%20')}"
    subprocess.run([
        "am","start",
        "-a","android.intent.action.VIEW",
        "-d",url
    ])

# ================= CONTADOR FACTURA =================

if not os.path.exists(CONTADOR):
    with open(CONTADOR,"w") as f:
        f.write("0")

with open(CONTADOR,"r+") as f:
    n = int(f.read()) + 1
    f.seek(0); f.write(str(n)); f.truncate()

FACTURA = f"FAC-{n:04d}"
FECHA = datetime.now()
FECHA_STR = FECHA.strftime("%Y-%m-%d")

print(f"\n=== {EMPRESA} | SERVICIO TÉCNICO ===\n")

# ================= CLIENTE =================

cliente   = pedir_texto("Cliente: ")
telefono  = pedir_texto("Teléfono (ej: 573001234567): ")
ciudad    = pedir_texto("Ciudad: ")
direccion = pedir_texto("Dirección: ")

# ================= MÁQUINA =================

modelo   = pedir_texto("Modelo de la máquina: ")
problema = pedir_texto("Problema reportado: ")
solucion = pedir_texto("Cómo se solucionó: ")
proceso  = pedir_texto("Proceso realizado: ")

# ================= TRABAJO =================

tipo = pedir_opcion(
    "\nTipo de trabajo:",
    {"1":"Diagnóstico","2":"Servicio","3":"Acción reparadora"}
)

mano_obra = pedir_precio(f"Precio mano de obra ({tipo}): ")

# ================= INSUMOS =================

insumos = []
total_insumos = 0.0

resp = input(
    "\nInsumos cambiados "
    "(coma) o 'no' / 'ninguno' / ENTER: "
).strip().lower()

if resp not in ("", "no", "ninguno"):
    nombres = [i.strip() for i in resp.split(",")]
    marcas = pedir_texto("Marcas (mismo orden): ").split(",")
    contador = pedir_texto("Contador de la máquina: ")
    precios = pedir_texto("Precios (mismo orden): ").split(",")

    if len(nombres)==len(marcas)==len(precios):
        for i in range(len(nombres)):
            precio = pedir_precio(f"Precio {nombres[i]}: ")
            insumos.append((nombres[i],marcas[i],contador,precio))
            total_insumos += precio

            archivo = f"{BASE_INSUMOS}/{slug(nombres[i])}.json"
            data = {"nombre":nombres[i],"total_cambios":0,"historial":[]}

            if os.path.exists(archivo):
                with open(archivo) as f:
                    data = json.load(f)

            data["total_cambios"] += 1
            data["historial"].append({
                "fecha":FECHA_STR,
                "cliente":cliente,
                "modelo":modelo,
                "marca":marcas[i],
                "contador":contador
            })

            with open(archivo,"w") as f:
                json.dump(data,f,indent=2,ensure_ascii=False)

# ================= TOTALES =================

TOTAL = mano_obra + total_insumos

# ================= CARPETAS =================

ruta = f"{BASE_CLIENTES}/{slug(cliente)}/{FECHA_STR}"
os.makedirs(f"{ruta}/pruebas_impresion", exist_ok=True)

print("\n📸 Tome fotos de las pruebas de impresión.")
input("ENTER para continuar...")

# ================= REPORTE WORD =================

doc = Document()

if os.path.exists(LOGO):
    doc.add_picture(LOGO, width=Inches(1.3))

doc.add_heading(EMPRESA,1)
doc.add_paragraph(f"Técnico: {TECNICO}")
doc.add_paragraph(f"Factura: {FACTURA} | Fecha: {FECHA_STR}")

doc.add_heading("Detalle",2)
doc.add_paragraph(f"Cliente: {cliente}")
doc.add_paragraph(f"Modelo: {modelo}")
doc.add_paragraph(f"Problema: {problema}")
doc.add_paragraph(f"Solución: {solucion}")
doc.add_paragraph(f"Proceso: {proceso}")
doc.add_paragraph(f"Tipo de trabajo: {tipo}")

if insumos:
    doc.add_heading("Insumos",2)
    for n,m,c,p in insumos:
        doc.add_paragraph(f"- {n} ({m}) | ${p}")

doc.add_heading("Costos",2)
doc.add_paragraph(f"Mano de obra: ${mano_obra}")
if insumos:
    doc.add_paragraph(f"Insumos: ${total_insumos}")
doc.add_paragraph(f"TOTAL: ${TOTAL}")

doc.save(f"{ruta}/{FACTURA}_reporte.docx")

# ================= FACTURA PDF (COMPACTA) =================

pdf = canvas.Canvas(f"{ruta}/{FACTURA}_factura.pdf", pagesize=A5)

y = 390
def linea(t):
    global y
    pdf.drawString(30,y,t)
    y -= 14

if os.path.exists(LOGO):
    pdf.drawImage(LOGO,30,420,width=80,height=40)

linea(EMPRESA)
linea(f"Técnico: {TECNICO}")
linea(f"Factura: {FACTURA}")
linea(f"Fecha: {FECHA_STR}")
linea("")
linea(f"Cliente: {cliente}")
linea(f"Trabajo: {tipo}")
linea("")
linea(f"Mano de obra: ${mano_obra}")

for n,m,c,p in insumos:
    linea(f"{n} ({m}) ${p}")

linea(f"TOTAL: ${TOTAL}")

pdf.showPage()
pdf.save()

print("\n✅ Servicio registrado correctamente")
print(f"📂 Carpeta: {ruta}")

# ================= WHATSAPP =================

op = input("\n¿Enviar por WhatsApp? (s/n): ").lower()

if op=="s":
    print(f"Número actual: {telefono}")
    if input("¿Cambiar número? (s/n): ").lower()=="s":
        telefono = pedir_texto("Nuevo número (ej: 573001234567): ")

    mensaje = (
        f"Hola, te envío el reporte y la factura "
        f"del servicio técnico realizado el {FECHA_STR}. "
        f"Quedo atento."
    )

    enviar_whatsapp(telefono, mensaje)

# ================= ABRIR CARPETA =================

subprocess.run([
    "am","start","--user","0",
    "-a","android.intent.action.VIEW",
    "-d",f"file://{ruta}"
])
